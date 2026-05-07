import os
import streamlit as st
from docx import Document
from pypdf import PdfReader

from langchain_core.documents import Document as LangchainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS


# ==================================================
# CONFIG
# ==================================================
st.set_page_config(
    page_title="MFU Grant Assistant using RAG",
    page_icon="📚",
    layout="wide"
)


# ==================================================
# CUSTOM CSS
# ==================================================
st.markdown("""
<style>
.stApp {
    background: linear-gradient(135deg, #f7f4ff 0%, #eef7ff 45%, #ffffff 100%);
}

.main-title {
    font-size: 42px;
    font-weight: 800;
    color: #4B2E83;
    margin-bottom: 5px;
}

.subtitle {
    font-size: 18px;
    color: #555;
    margin-bottom: 25px;
}

.hero-card {
    background: white;
    padding: 28px;
    border-radius: 22px;
    box-shadow: 0 8px 24px rgba(75, 46, 131, 0.12);
    border-left: 8px solid #6C4AB6;
}

.info-card {
    background: #ffffff;
    padding: 20px;
    border-radius: 18px;
    box-shadow: 0 5px 18px rgba(0,0,0,0.08);
    height: 100%;
}

.metric-title {
    color: #6C4AB6;
    font-size: 18px;
    font-weight: 700;
}

.answer-box {
    background: #ffffff;
    padding: 24px;
    border-radius: 18px;
    border: 1px solid #e5ddff;
    box-shadow: 0 5px 18px rgba(0,0,0,0.06);
}

.source-box {
    background: #f4f0ff;
    padding: 12px 16px;
    border-radius: 12px;
    margin-bottom: 8px;
    color: #4B2E83;
    font-weight: 600;
}

.footer {
    text-align: center;
    color: #777;
    font-size: 14px;
}
</style>
""", unsafe_allow_html=True)


# ==================================================
# GROUP INFO
# ==================================================
GROUP_NO = "BDA_Project2_GroupNo1"

GROUP_MEMBERS = [
    {"student_id": "6631501120", "name": "Achira Lueablae"},
    {"student_id": "6631501121", "name": "Aphiwat Chioewvijit"},
    {"student_id": "6631501122", "name": "Alicha Chanatnawa"},
    {"student_id": "6631501125", "name": "Araya Mahima"},
    {"student_id": "6631501130", "name": "Catarina Magdaleno Roquette"},
]


DATASET_DIR = "dataset"


# ==================================================
# SIDEBAR
# ==================================================
with st.sidebar:
    st.markdown("## 📘 Project Info")
    st.markdown(f"### **{GROUP_NO}**")

    st.divider()

    st.markdown("### 👥 Group Members")
    for member in GROUP_MEMBERS:
        st.markdown(f"**{member['student_id']}**  \n{member['name']}")

    st.divider()

    st.markdown("### 💡 Example Questions")
    st.markdown("""
    - ทุนประเภท 1 กับ 2 ต่างกันอย่างไร  
    - ได้เงินสนับสนุนเมื่อไหร่  
    - ถ้าไม่ผ่านระดับดีต้องทำอย่างไร  
    - ลิขสิทธิ์กี่ปี  
    - ผู้เขียนได้ค่าลิขสิทธิ์กี่เปอร์เซ็นต์  
    """)


# ==================================================
# FILE READERS
# ==================================================
def read_docx(file_path):
    try:
        doc = Document(file_path)
        text = []

        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())

        return "\n".join(text)

    except Exception as e:
        return f"Error reading DOCX: {e}"


def read_pdf(file_path):
    text = []

    try:
        reader = PdfReader(file_path)

        for page in reader.pages:
            page_text = page.extract_text()

            if page_text:
                text.append(page_text)

    except Exception as e:
        text.append(f"Error reading PDF: {e}")

    return "\n".join(text)


# ==================================================
# IMAGE SUMMARY
# ==================================================
IMAGE_KNOWLEDGE = """
ทุนประเภทที่ 1:
- เพื่อขอตำแหน่งทางวิชาการ
- มหาวิทยาลัยสนับสนุน 30,000 บาท
- ผ่านสำนักวิชา
- ผู้ทรงคุณวุฒิ 2 คน
- ต้องผ่านระดับดีขึ้นไป
- ได้ค่าลิขสิทธิ์ 30%

ทุนประเภทที่ 2:
- eBook เพื่อการจำหน่าย
- ผู้เขียนส่งเอง
- ผู้เขียนจ่ายค่าผู้ทรงคุณวุฒิ
- ต้องผ่านระดับดีขึ้นไป
- MLii จัดทำ eBook
- ได้รายได้ 80% หลังหักค่าใช้จ่าย
- มอบลิขสิทธิ์ 5 ปี
"""


# ==================================================
# LOAD DOCUMENTS
# ==================================================
def load_documents():
    documents = []

    if not os.path.exists(DATASET_DIR):
        return documents

    for root, dirs, files in os.walk(DATASET_DIR):

        for file in files:

            path = os.path.join(root, file)

            content = ""

            if file.lower().endswith(".docx"):
                content = read_docx(path)

            elif file.lower().endswith(".pdf"):
                content = read_pdf(path)

            else:
                continue

            if content.strip():
                documents.append(
                    LangchainDocument(
                        page_content=content,
                        metadata={"source": file}
                    )
                )

    documents.append(
        LangchainDocument(
            page_content=IMAGE_KNOWLEDGE,
            metadata={"source": "ทุนประเภทที่1-2 Infographic Summary"}
        )
    )

    return documents


# ==================================================
# VECTORSTORE
# ==================================================
@st.cache_resource
def create_vectorstore():

    documents = load_documents()

    if len(documents) == 0:
        return None

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150
    )

    chunks = splitter.split_documents(documents)

    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
    )

    vectorstore = FAISS.from_documents(chunks, embeddings)

    return vectorstore


# ==================================================
# ANSWER FUNCTION
# ==================================================
def get_answer(question, vectorstore):

    docs = vectorstore.similarity_search(question, k=4)

    if not docs:
        return "ไม่พบข้อมูลในเอกสาร", []

    answer = "### สรุปคำตอบจาก Dataset\n"

    sources = []

    for i, doc in enumerate(docs, start=1):

        answer += f"\n\nข้อมูลอ้างอิง {i}:\n"
        answer += doc.page_content[:800]

        source = doc.metadata.get("source", "Unknown")
        sources.append(source)

    return answer, sources


# ==================================================
# MAIN UI
# ==================================================
st.markdown('<div class="main-title">📚 MFU Grant Assistant using RAG</div>', unsafe_allow_html=True)

st.markdown(
    '<div class="subtitle">AI Chatbot สำหรับตอบคำถามเกี่ยวกับทุนตำรา หนังสือ eBook และขั้นตอนการขอรับทุน</div>',
    unsafe_allow_html=True
)

st.markdown("""
<div class="hero-card">
<h3>🎯 Project Overview</h3>
<p>
ระบบนี้ใช้แนวคิด <b>Retrieval-Augmented Generation (RAG)</b> 
เพื่อค้นคืนข้อมูลจากเอกสาร Dataset ของทุนตำรา MFU แล้วนำข้อมูลที่เกี่ยวข้องมาตอบคำถามผู้ใช้
</p>
</div>
""", unsafe_allow_html=True)

st.write("")

col1, col2, col3 = st.columns(3)

with col1:
    st.markdown("""
    <div class="info-card">
        <div class="metric-title">📄 Dataset</div>
        <p>PDF, DOCX และ Infographic เกี่ยวกับทุนตำรา MFU</p>
    </div>
    """, unsafe_allow_html=True)

with col2:
    st.markdown("""
    <div class="info-card">
        <div class="metric-title">🧠 AI Concept</div>
        <p>Document Loading, Text Splitting, Embedding และ FAISS Retrieval</p>
    </div>
    """, unsafe_allow_html=True)

with col3:
    st.markdown("""
    <div class="info-card">
        <div class="metric-title">🚀 Deployment</div>
        <p>พัฒนา Web Application ด้วย Streamlit Cloud</p>
    </div>
    """, unsafe_allow_html=True)

st.write("")
st.divider()

vectorstore = create_vectorstore()

if vectorstore is None:
    st.error("ไม่พบ dataset กรุณาตรวจสอบโฟลเดอร์ dataset")
    st.stop()

st.markdown("## 🔍 Ask from Dataset")

example_questions = [
    "ทุนประเภทที่ 1 กับประเภทที่ 2 ต่างกันอย่างไร",
    "ถ้าไม่ผ่านระดับดีต้องทำอย่างไร",
    "ลิขสิทธิ์กี่ปี",
    "ผู้เขียนได้รับค่าลิขสิทธิ์กี่เปอร์เซ็นต์"
]

selected_example = st.selectbox(
    "เลือกคำถามตัวอย่าง หรือพิมพ์คำถามเองด้านล่าง",
    [""] + example_questions
)

question = st.text_input(
    "กรอกคำถามของคุณ:",
    value=selected_example,
    placeholder="เช่น ทุนประเภทที่ 1 กับประเภทที่ 2 ต่างกันอย่างไร?"
)

if question:
    with st.spinner("กำลังวิเคราะห์ข้อมูลจาก Dataset..."):
        answer, sources = get_answer(question, vectorstore)

    st.markdown("## ✅ Answer")
    st.markdown(f'<div class="answer-box">{answer}</div>', unsafe_allow_html=True)

    st.markdown("## 📌 Sources Used")
    unique_sources = list(set(sources))

    for src in unique_sources:
        st.markdown(f'<div class="source-box">📎 {src}</div>', unsafe_allow_html=True)

st.divider()

st.markdown(
    '<div class="footer">Developed using Vibecode + RAG + Streamlit | Business Data Analytics Project</div>',
    unsafe_allow_html=True
)
